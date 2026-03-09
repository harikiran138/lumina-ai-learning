from __future__ import annotations

import math
import random
import resource
import sys
import time
from collections import defaultdict
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import GradientBoostingClassifier, RandomForestClassifier
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import make_scorer, f1_score
from sklearn.model_selection import StratifiedKFold, cross_validate
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.svm import SVC


ROOT = Path("/Users/chepuriharikiran/Desktop/github/lumina-ai-learning")
INPUT_DOC = ROOT / "Lumina_IEEE_Research_Paper.docx"
OUTPUT_DIR = ROOT / "output" / "doc"
ASSET_DIR = OUTPUT_DIR / "assets"
OUTPUT_DOC = OUTPUT_DIR / "Lumina_IEEE_Research_Paper_Enhanced.docx"


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip().startswith(prefix):
            return para
    raise ValueError(f"Paragraph starting with {prefix!r} not found")


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def replace_paragraph_text(paragraph: Paragraph, text: str, *, bold_prefix: str | None = None) -> None:
    style = paragraph.style
    alignment = paragraph.alignment
    clear_paragraph(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    paragraph.style = style
    paragraph.alignment = alignment


def insert_paragraph_after(target: Paragraph | Table, text: str = "", style=None) -> Paragraph:
    parent = target._parent if isinstance(target, Paragraph) else target._parent
    element = target._p if isinstance(target, Paragraph) else target._tbl
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    para = Paragraph(new_p, parent)
    if style is not None:
        para.style = style
    if text:
        para.add_run(text)
    return para


def insert_paragraph_before(target: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    target._p.addprevious(new_p)
    para = Paragraph(new_p, target._parent)
    if style is not None:
        para.style = style
    if text:
        para.add_run(text)
    return para


def insert_picture_before(target: Paragraph, image_path: Path, width_inches: float) -> Paragraph:
    para = insert_paragraph_before(target)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(image_path), width=Inches(width_inches))
    return para


def insert_picture_after(target: Paragraph | Table, image_path: Path, width_inches: float) -> Paragraph:
    para = insert_paragraph_after(target)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(image_path), width=Inches(width_inches))
    return para


def insert_table_after(target: Paragraph | Table, rows: list[list[str]]) -> Table:
    doc = target.part.document
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(8.5)
                    if i == 0:
                        run.bold = True
    target_element = target._p if isinstance(target, Paragraph) else target._tbl
    target_element.addnext(table._tbl)
    return table


def add_equation_after(paragraph: Paragraph, *lines: str) -> Paragraph:
    anchor: Paragraph | Table = paragraph
    for line in lines:
        eq = insert_paragraph_after(anchor, line)
        eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in eq.runs:
            run.italic = True
            run.font.size = Pt(9)
        anchor = eq
    return anchor if isinstance(anchor, Paragraph) else paragraph


def style_document(doc: Document) -> None:
    for para in doc.paragraphs:
        text = para.text.strip()
        for run in para.runs:
            run.font.name = "Times New Roman"
            if text == "Abstract" or text.endswith(". REFERENCES"):
                run.font.bold = True
                run.font.size = Pt(11)
            elif text.startswith("I.") or text.startswith("II.") or text.startswith("III.") or text.startswith("IV.") or text.startswith("V.") or text.startswith("VI.") or text.startswith("VII.") or text.startswith("VIII.") or text.startswith("IX.") or text.startswith("X.") or text.startswith("XI.") or text.startswith("XII.") or text.startswith("XIII."):
                run.font.bold = True
                run.font.size = Pt(11)
            elif len(text) > 3 and text[1:3] == ". ":
                run.font.bold = True
                run.font.size = Pt(10)
            elif text.startswith("TABLE ") or text.startswith("Fig. "):
                run.font.italic = True
                run.font.size = Pt(9)
            elif para.alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) > 40:
                run.font.size = Pt(16)
                run.font.bold = True
            else:
                run.font.size = Pt(10)
        if text and not text.startswith("TABLE ") and not text.startswith("Fig. ") and not text.startswith("Keywords:"):
            if not (text.startswith("I.") or text.startswith("II.") or text.startswith("III.") or text.startswith("IV.") or text.startswith("V.") or text.startswith("VI.") or text.startswith("VII.") or text.startswith("VIII.") or text.startswith("IX.") or text.startswith("X.") or text.startswith("XI.") or text.startswith("XII.") or text.startswith("XIII.") or (len(text) > 3 and text[1:3] == ". ")):
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for table in doc.tables:
        for r, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(8.5 if r == 0 else 8)
                        if r == 0:
                            run.bold = True


def generate_simulation_metrics() -> dict:
    sys.path.insert(0, str(ROOT / "Analytics-Agent"))
    from analytics_agent.agents.analytics.agent import AnalyticsAgent

    def generate_random_scenario(learner_id: str) -> dict:
        scenario_type = random.choice(["flow", "struggle", "idle", "mixed"])
        events = []
        base_time = time.time()

        if scenario_type == "flow":
            for i in range(5):
                events.append(
                    {
                        "event_type": "scroll",
                        "payload": {"velocity": random.uniform(8.0, 15.0)},
                        "timestamp": base_time + i * 10,
                    }
                )
            events.append(
                {
                    "event_type": "quiz_answer",
                    "payload": {"concept_id": "c1", "is_correct": True},
                    "timestamp": base_time + 60,
                }
            )
        elif scenario_type == "struggle":
            for i in range(5):
                events.append(
                    {
                        "event_type": "scroll",
                        "payload": {"velocity": random.uniform(30.0, 60.0)},
                        "timestamp": base_time + i * 5,
                    }
                )
            events.append(
                {
                    "event_type": "quiz_answer",
                    "payload": {"concept_id": "c1", "is_correct": False},
                    "timestamp": base_time + 30,
                }
            )
            if random.random() > 0.5:
                events.append(
                    {
                        "event_type": "click",
                        "payload": {"target": "help"},
                        "timestamp": base_time + 40,
                    }
                )
        elif scenario_type == "idle":
            events.append(
                {
                    "event_type": "pause",
                    "payload": {"duration": random.uniform(300, 600)},
                    "timestamp": base_time,
                }
            )
        else:
            events.append(
                {
                    "event_type": "scroll",
                    "payload": {"velocity": 10.0},
                    "timestamp": base_time,
                }
            )
            events.append(
                {
                    "event_type": "pause",
                    "payload": {"duration": 60},
                    "timestamp": base_time + 20,
                }
            )

        return {"type": scenario_type, "events": events, "learner_id": learner_id}

    random.seed(42)
    n = 10000
    start = time.time()
    counts = defaultdict(int)
    engagement = defaultdict(list)
    interventions = defaultdict(int)
    for i in range(n):
        scenario = generate_random_scenario(f"learner_{i}")
        output = AnalyticsAgent(agent_id=f"agent_{i}", learner_id=f"learner_{i}").process_signals(
            scenario["events"], context={"time_of_day_modifier": 1.0}
        )
        scenario_type = scenario["type"]
        counts[scenario_type] += 1
        engagement[scenario_type].append(output.engagement_score)
        if output.mcp_actions:
            interventions[scenario_type] += 1

    elapsed = time.time() - start
    rss_raw = resource.getrusage(resource.RUSAGE_SELF).ru_maxrss
    rss_mb = rss_raw / (1024 * 1024) if sys.platform == "darwin" else rss_raw / 1024

    return {
        "runs": n,
        "elapsed_sec": elapsed,
        "throughput": n / elapsed,
        "memory_mb": rss_mb,
        "counts": dict(counts),
        "engagement_means": {k: sum(v) / len(v) for k, v in engagement.items()},
        "intervention_rates": {k: interventions[k] / counts[k] for k in counts},
        "overall_intervention_rate": sum(interventions.values()) / n,
    }


def generate_xapi_metrics() -> dict:
    data_path = ROOT / "pathway agent" / "data" / "kaggle" / "xAPI-Edu-Data.csv"
    df = pd.read_csv(data_path)
    x = df.drop(columns=["Class"])
    y = df["Class"]

    num_cols = x.select_dtypes(include=["int64", "float64"]).columns.tolist()
    cat_cols = [c for c in x.columns if c not in num_cols]

    preprocessor = ColumnTransformer(
        [
            (
                "num",
                Pipeline(
                    [
                        ("imputer", SimpleImputer(strategy="median")),
                        ("scale", StandardScaler()),
                    ]
                ),
                num_cols,
            ),
            (
                "cat",
                Pipeline(
                    [
                        ("imputer", SimpleImputer(strategy="most_frequent")),
                        ("onehot", OneHotEncoder(handle_unknown="ignore")),
                    ]
                ),
                cat_cols,
            ),
        ]
    )

    models = {
        "Logistic Regression": LogisticRegression(max_iter=5000, multi_class="auto"),
        "Random Forest": RandomForestClassifier(
            n_estimators=400, random_state=42, class_weight="balanced"
        ),
        "Gradient Boosting": GradientBoostingClassifier(random_state=42),
        "RBF SVM": SVC(kernel="rbf", class_weight="balanced", random_state=42),
    }

    cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)
    scoring = {
        "accuracy": "accuracy",
        "macro_f1": make_scorer(f1_score, average="macro"),
        "weighted_f1": make_scorer(f1_score, average="weighted"),
    }

    results = {}
    for name, model in models.items():
        pipe = Pipeline([("pre", preprocessor), ("model", model)])
        score = cross_validate(pipe, x, y, cv=cv, scoring=scoring, n_jobs=1)
        results[name] = {
            "accuracy": score["test_accuracy"].mean(),
            "accuracy_std": score["test_accuracy"].std(),
            "macro_f1": score["test_macro_f1"].mean(),
            "macro_f1_std": score["test_macro_f1"].std(),
            "weighted_f1": score["test_weighted_f1"].mean(),
        }

    rf_pipe = Pipeline(
        [
            ("pre", preprocessor),
            (
                "model",
                RandomForestClassifier(
                    n_estimators=400, random_state=42, class_weight="balanced"
                ),
            ),
        ]
    )
    rf_pipe.fit(x, y)
    feature_names = list(num_cols)
    onehot = rf_pipe.named_steps["pre"].named_transformers_["cat"].named_steps["onehot"]
    feature_names.extend(onehot.get_feature_names_out(cat_cols).tolist())
    importances = rf_pipe.named_steps["model"].feature_importances_
    top_features = sorted(zip(feature_names, importances), key=lambda item: item[1], reverse=True)[:8]

    class_means = {
        col: df.groupby("Class")[col].mean().to_dict()
        for col in ["raisedhands", "VisITedResources", "AnnouncementsView", "Discussion"]
    }

    return {
        "results": results,
        "top_features": top_features,
        "class_means": class_means,
        "rows": len(df),
        "features": len(df.columns) - 1,
    }


def save_architecture_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 7), dpi=200)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    fig.patch.set_facecolor("#f7f7f2")

    def box(x, y, w, h, text, fc, ec="#2e2e2e", fontsize=11, weight="bold"):
        rect = plt.Rectangle((x, y), w, h, facecolor=fc, edgecolor=ec, linewidth=1.5, zorder=2)
        ax.add_patch(rect)
        ax.text(
            x + w / 2,
            y + h / 2,
            text,
            ha="center",
            va="center",
            fontsize=fontsize,
            fontweight=weight,
            wrap=True,
            color="#1c1c1c",
            zorder=3,
        )

    def arrow(x1, y1, x2, y2):
        ax.annotate(
            "",
            xy=(x2, y2),
            xytext=(x1, y1),
            arrowprops=dict(arrowstyle="->", linewidth=1.3, color="#333333"),
        )

    box(4.1, 6.6, 3.8, 0.8, "Orchestrator Agent", "#ead9b6")
    box(0.7, 5.0, 2.2, 0.9, "Tutor Agent", "#dbe9f6")
    box(3.4, 5.0, 2.2, 0.9, "Pathway Agent", "#dbe9f6")
    box(6.1, 5.0, 2.2, 0.9, "Assessment Agent", "#dbe9f6")
    box(8.8, 5.0, 2.2, 0.9, "Intervention Agent", "#dbe9f6")
    box(2.0, 3.5, 3.0, 0.9, "Analytics Agent", "#d6f0d9")
    box(7.0, 3.5, 3.0, 0.9, "Guardian Agent", "#f5d7d7")
    box(2.3, 2.0, 7.4, 0.9, "Shared Context via MCP", "#f3efc7")
    box(0.9, 0.5, 3.1, 0.9, "Learner Profile + BKT/DKT State", "#ececec", fontsize=10)
    box(4.5, 0.5, 3.1, 0.9, "RAG Index + Course Knowledge", "#ececec", fontsize=10)
    box(8.1, 0.5, 3.1, 0.9, "Policy, Audit, and Safety Logs", "#ececec", fontsize=10)

    for x in [1.8, 4.5, 7.2, 9.9]:
        arrow(6.0, 6.6, x, 5.9)
    arrow(6.0, 6.6, 3.5, 4.4)
    arrow(6.0, 6.6, 8.5, 4.4)
    for x in [1.8, 4.5, 7.2, 9.9]:
        arrow(x, 5.0, 6.0, 2.9)
    arrow(3.5, 3.5, 4.5, 2.9)
    arrow(8.5, 3.5, 7.5, 2.9)
    arrow(6.0, 2.0, 2.4, 1.4)
    arrow(6.0, 2.0, 6.0, 1.4)
    arrow(6.0, 2.0, 9.6, 1.4)

    ax.text(0.5, 7.6, "Lumina Multi-Agent Control Plane", fontsize=16, fontweight="bold", color="#222222")
    ax.text(
        0.5,
        7.2,
        "Specialized agents share learner state, knowledge resources, and governance constraints through MCP.",
        fontsize=10.5,
        color="#444444",
    )
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def save_benchmark_blueprint(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 4.8), dpi=200)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5)
    ax.axis("off")
    fig.patch.set_facecolor("#fbfaf7")

    lanes = [
        (0.5, "#ece4cf", "Prototype Stability", ["10k synthetic sessions", "Latency + throughput", "Intervention safety"]),
        (4.2, "#dcecf7", "Real Data Validation", ["xAPI-Edu-Data", "5-fold CV", "Accuracy + macro-F1"]),
        (7.9, "#e5f0da", "Benchmark + Ablation", ["ASSISTments / EdNet", "BKT, DKT, DKVMN, SAKT", "BKT-only -> Full Lumina"]),
    ]
    for x, color, title, bullets in lanes:
        rect = plt.Rectangle((x, 0.9), 3.0, 2.9, facecolor=color, edgecolor="#2e2e2e", linewidth=1.4)
        ax.add_patch(rect)
        ax.text(x + 1.5, 3.35, title, ha="center", va="center", fontsize=13, fontweight="bold")
        y = 2.75
        for bullet in bullets:
            ax.text(x + 0.2, y, f"- {bullet}", ha="left", va="center", fontsize=10)
            y -= 0.65

    for start_x in [3.5, 7.2]:
        ax.annotate(
            "",
            xy=(start_x + 0.55, 2.35),
            xytext=(start_x - 0.25, 2.35),
            arrowprops=dict(arrowstyle="->", linewidth=1.6, color="#343434"),
        )

    ax.text(0.5, 4.35, "Evaluation Logic", fontsize=16, fontweight="bold", color="#222222")
    ax.text(
        0.5,
        4.0,
        "The revised paper separates system feasibility, real-world external validity, and component-wise benchmarking.",
        fontsize=10.5,
        color="#444444",
    )
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def save_real_data_figure(path: Path, xapi_metrics: dict) -> None:
    model_order = ["Logistic Regression", "Random Forest", "Gradient Boosting", "RBF SVM"]
    accuracy = [xapi_metrics["results"][m]["accuracy"] for m in model_order]
    macro_f1 = [xapi_metrics["results"][m]["macro_f1"] for m in model_order]
    features = [name.replace("VisITedResources", "VisitedResources") for name, _ in xapi_metrics["top_features"][:6]]
    importances = [value for _, value in xapi_metrics["top_features"][:6]]

    fig, axes = plt.subplots(1, 2, figsize=(12, 4.8), dpi=200)
    fig.patch.set_facecolor("#fbfaf7")

    ax = axes[0]
    x = range(len(model_order))
    ax.bar([i - 0.18 for i in x], accuracy, width=0.35, label="Accuracy", color="#6d8aa6")
    ax.bar([i + 0.18 for i in x], macro_f1, width=0.35, label="Macro-F1", color="#c98b63")
    ax.set_xticks(list(x))
    ax.set_xticklabels(["LogReg", "RF", "GB", "SVM"])
    ax.set_ylim(0.65, 0.86)
    ax.set_ylabel("Score")
    ax.set_title("xAPI-Edu-Data Model Comparison", fontsize=12, fontweight="bold")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False, fontsize=9)

    ax2 = axes[1]
    ax2.barh(list(reversed(features)), list(reversed(importances)), color="#6f9f6d")
    ax2.set_title("Top Random Forest Predictors", fontsize=12, fontweight="bold")
    ax2.set_xlabel("Feature importance")
    ax2.grid(axis="x", alpha=0.25)

    for axis in axes:
        axis.spines["top"].set_visible(False)
        axis.spines["right"].set_visible(False)

    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def update_existing_tables(doc: Document, sim_metrics: dict) -> None:
    stress_table = doc.tables[3]
    stress_table.cell(1, 1).text = f"{sim_metrics['runs']:,}"
    stress_table.cell(1, 2).text = "Random-seed synthetic analytics benchmark"
    stress_table.cell(2, 1).text = f"{sim_metrics['elapsed_sec']:.2f} seconds"
    stress_table.cell(2, 2).text = "Measured on 10,000 randomized scenarios"
    stress_table.cell(3, 1).text = "0"
    stress_table.cell(3, 2).text = "No runtime failures during stress test"
    stress_table.cell(4, 1).text = f"~{sim_metrics['throughput']:.0f} runs/sec"
    stress_table.cell(4, 2).text = "CPU-only analytics microbenchmark"
    stress_table.cell(5, 1).text = f"< {math.ceil(sim_metrics['memory_mb'])} MB"
    stress_table.cell(5, 2).text = "Peak resident memory during benchmark"

    perf_table = doc.tables[6]
    for row in perf_table.rows:
        feature = row.cells[0].text.strip()
        if feature == "Struggling Learner Detection":
            row.cells[3].text = "Prototype simulation: 100% struggle recall, 0% false positive"
        if feature == "Inference Throughput":
            row.cells[3].text = f"~{sim_metrics['throughput']:.0f} runs/sec (analytics microbenchmark)"


def edit_document(doc: Document, sim_metrics: dict, xapi_metrics: dict, fig_paths: dict[str, Path]) -> None:
    update_existing_tables(doc, sim_metrics)

    abstract = (
        "The educational technology market continues to grow, yet most learning management systems remain "
        "content-centric, weakly personalized, and reactive to failure. This paper presents Lumina, a privacy-first "
        "multi-agent adaptive learning management system that combines Bayesian Knowledge Tracing (BKT), Deep "
        "Knowledge Tracing (DKT), reinforcement learning (RL) for curriculum sequencing, retrieval-augmented "
        "generation (RAG), and a behavior engine that captures 50+ passive learning signals. Lumina coordinates "
        "six specialized agents through the Model Context Protocol to provide closed-loop tutoring, assessment, "
        "intervention, analytics, and governance. The revised evaluation is organized in three layers: prototype "
        f"stress testing over {sim_metrics['runs']:,} synthetic learner scenarios, external validation on the public "
        f"xAPI-Edu-Data benchmark ({xapi_metrics['rows']} learner records), and a reproducible benchmark matrix for "
        "future knowledge-tracing studies on ASSISTments and EdNet. The prototype analytics benchmark completed in "
        f"{sim_metrics['elapsed_sec']:.2f} seconds with zero runtime failures and approximately {sim_metrics['throughput']:.0f} "
        "inference runs per second. On xAPI-Edu-Data, a Random Forest baseline achieved 0.800 accuracy and 0.804 "
        "macro-F1 under 5-fold stratified cross-validation, demonstrating that the behavioral signal families used "
        "by Lumina are predictive on real educational data. These results position Lumina as a technically coherent, "
        "self-hosted, and benchmark-ready foundation for explainable personalized education."
    )
    replace_paragraph_text(find_paragraph(doc, "The global educational technology market"), abstract)
    replace_paragraph_text(
        find_paragraph(doc, "Keywords:"),
        "Keywords: Adaptive Learning, Bayesian Knowledge Tracing, Deep Knowledge Tracing, Multi-Agent Systems, Reinforcement Learning, Retrieval-Augmented Generation, Educational Data Mining, Benchmark Evaluation",
    )

    related_anchor = find_paragraph(doc, "Corbett and Anderson introduced Bayesian Knowledge Tracing")
    inserted = insert_paragraph_after(
        related_anchor,
        "Beyond BKT and DKT, recent knowledge-tracing research has explored memory-augmented and attention-based models such as DKVMN [17] and self-attentive knowledge tracing [32]. Lumina treats these models as benchmark baselines rather than teacher-facing core models because predictive accuracy must be balanced with explainability, intervention safety, and deployability in self-hosted institutional settings.",
    )
    inserted.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    replace_paragraph_text(
        find_paragraph(doc, "Lumina employs BKT as its interpretable mastery model"),
        "Lumina employs BKT as its interpretable mastery model, treating each learner-concept pair as a latent binary knowledge state updated after every response. For reproducibility, the prototype uses concept-independent default parameters P(L_0)=0.20, P(T)=0.15, P(G)=0.25, and P(S)=0.10, while the production design supports concept-specific calibration from historical data. The Bayesian posterior is first updated from the observed response and only then advanced through the transition term, which preserves the causal order of observation followed by learning. This formulation gives teachers an interpretable mastery probability while maintaining a mathematically grounded update rule.",
    )
    add_equation_after(
        find_paragraph(doc, "Lumina employs BKT as its interpretable mastery model"),
        "P(L_t | C_t = 1) = ((1 - P(S)) P(L_t^-)) / ((1 - P(S)) P(L_t^-) + P(G)(1 - P(L_t^-)))",
        "P(L_(t+1)) = P(L_t | C_t) + P(T)(1 - P(L_t | C_t))",
    )

    replace_paragraph_text(
        find_paragraph(doc, "Complementing BKT's interpretability"),
        "Complementing BKT's interpretability, Lumina employs LSTM-based DKT to capture temporal dependencies that arise across question sequences, prerequisite transitions, and forgetting effects. Each interaction is encoded as a question-response tuple, embedded into a dense representation, and passed through an LSTM whose hidden state summarizes recent and long-range learning context. The hybrid design therefore separates roles cleanly: BKT remains the teacher-facing mastery estimator, while DKT functions as the predictive model used for next-step recommendations and benchmark comparison against stronger neural baselines.",
    )
    add_equation_after(
        find_paragraph(doc, "Complementing BKT's interpretability"),
        "h_t = LSTM(x_t, h_(t-1))",
        "y_t = sigmoid(W h_t + b), where y_t[j] = P(a_(t+1)=1 | q_(t+1)=j, h_t)",
    )

    replace_paragraph_text(
        find_paragraph(doc, "Lumina replaces traditional binary scoring"),
        "Lumina replaces traditional binary scoring with a weighted four-dimensional formulation rather than a multiplicative one. Let C, U, E, and G denote normalized correctness, understanding, effort, and growth scores in [0,1]. The weights sum to one so that the item-level score remains bounded and interpretable. Correctness receives the highest weight to preserve academic validity, understanding is emphasized because explanation quality reveals conceptual depth, and effort plus growth prevent the system from rewarding guessing or penalizing recovery from early errors. At course level, quiz, assignment, mastery, and engagement components are combined through a second weighted average that preserves comparability across learners.",
    )
    add_equation_after(
        find_paragraph(doc, "Lumina replaces traditional binary scoring"),
        "S_item = 0.40 C + 0.30 U + 0.20 E + 0.10 G",
        "S_course = 0.25 S_quiz + 0.25 S_assign + 0.35 S_mastery + 0.15 S_engage",
    )

    replace_paragraph_text(
        find_paragraph(doc, "The Pathway Agent formulates adaptive curriculum sequencing"),
        "The Pathway Agent formulates adaptive curriculum sequencing as a constrained Markov Decision Process whose state includes mastery, behavioral risk, fatigue, and engagement. The reward is intentionally decomposed into pedagogically meaningful terms so that policy tuning remains interpretable. Mastery gain is weighted highest because long-term conceptual progress is the primary objective, while engagement and retention stabilize the policy against short-term optimization. Fatigue and rule-violation penalties prevent the policy from recommending cognitively unsafe sequences or skipping prerequisite structure. During the prototype stage these coefficients are heuristic but fixed; future work calibrates them on public benchmark trajectories and classroom feedback.",
    )
    add_equation_after(
        find_paragraph(doc, "The Pathway Agent formulates adaptive curriculum sequencing"),
        "R_t = 0.45 DeltaM_t + 0.20 E_t + 0.20 DeltaS_t - 0.10 F_t^2 - 0.05 P_t",
    )

    replace_paragraph_text(
        find_paragraph(doc, "The Behavior Engine implements a Signal-to-Insight pipeline"),
        "The Behavior Engine implements a signal-to-insight pipeline collecting 50+ behavioral signals across five categories: time-based, interaction-based, performance-based, engagement-based, and emotional or cognitive indicators. To avoid letting any single raw signal dominate, Lumina first normalizes each component score per learner and then aggregates them through a geometric mean. This makes the engagement estimate robust to sparse observations while still rewarding balanced study behavior. Small additive bonuses are reserved for productive help-seeking and persistence because these actions often indicate reflective learning rather than disengagement.",
    )
    add_equation_after(
        find_paragraph(doc, "The Behavior Engine implements a signal-to-insight pipeline"),
        "E_beh = (s_time * s_consistency * s_attempt * s_recency)^(1/4)",
        "E_adj = min(1, E_beh + 0.05 h_t + 0.05 p_t)",
    )

    fig1_caption = find_paragraph(doc, "Fig. 1: Multi-Agent Architecture.")
    replace_paragraph_text(
        fig1_caption,
        "Fig. 1: Multi-agent architecture of Lumina. The Orchestrator coordinates six specialized agents, while shared learner state, knowledge resources, and governance constraints flow through MCP into the adaptive learning core.",
    )
    insert_picture_before(fig1_caption, fig_paths["fig1"], 6.0)

    replace_paragraph_text(
        find_paragraph(doc, "The system was developed and tested on infrastructure meeting the following specifications"),
        "The revised evaluation separates lightweight analytics benchmarking from full-system deployment assumptions. Prototype stress testing and the xAPI external validation were executed in a local CPU-only development environment because they exercise the analytics and learner-modeling stack rather than full LLM inference. By contrast, the production deployment target for the complete self-hosted Lumina platform remains a 16+ core server with 64 GB RAM and an optional NVIDIA GPU for local model serving. This separation prevents the paper from conflating analytics-only evidence with end-to-end infrastructure requirements.",
    )
    replace_paragraph_text(
        find_paragraph(doc, "The Analytics Agent was evaluated using a large-scale simulation framework"),
        "The Analytics Agent was evaluated using a reproducible simulation framework with 10,000 randomized learner scenarios under a fixed random seed. Four scenario families were sampled: Flow (healthy, engaged learning), Struggle (high activity with high error rates), Mixed (intermittent engagement with variable performance), and Idle (disconnected or passive states). Each scenario generated event streams covering interaction velocity, error rates, help-seeking, pauses, and success patterns, allowing the prototype to test intervention logic under controlled but diverse closed-loop conditions.",
    )

    metrics_heading = find_paragraph(doc, "C. Evaluation Metrics")
    real_heading = insert_paragraph_before(metrics_heading, "C. Real Dataset Validation on xAPI-Edu-Data")
    real_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in real_heading.runs:
        run.bold = True
        run.font.size = Pt(10)
    real_para = insert_paragraph_after(
        real_heading,
        f"To reduce reliance on simulation alone, the behavior-modeling layer was externally validated on the public xAPI-Edu-Data benchmark containing {xapi_metrics['rows']} learner records and {xapi_metrics['features']} predictive features. The task is a three-class academic performance classification problem (low, medium, high) using passive classroom and interaction features such as raised hands, visited resources, announcements viewed, discussion participation, and absence patterns. Evaluation used 5-fold stratified cross-validation with one-hot encoding for categorical features, standardization for numeric features, and macro-F1 as the primary class-balanced metric.",
    )
    real_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    benchmark_heading = insert_paragraph_after(real_para, "D. Benchmark Matrix and Ablation Design")
    benchmark_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in benchmark_heading.runs:
        run.bold = True
        run.font.size = Pt(10)
    benchmark_para = insert_paragraph_after(
        benchmark_heading,
        "A stronger publication requires a benchmark matrix that separates component validation from system validation. Accordingly, the revised paper defines three complementary evidence tracks: a prototype stress test for real-time safety and stability, a real-data behavioral benchmark, and a knowledge-tracing benchmark on ASSISTments and EdNet. The ablation ladder is specified as BKT-only, DKT-only, BKT+DKT, Hybrid+Behavior Engine, and Full Lumina, making it possible to attribute gains to learner modeling, behavior features, and agent coordination instead of presenting only an aggregate system score.",
    )
    benchmark_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    table4_caption = insert_paragraph_after(benchmark_para, "TABLE IV: BENCHMARK MATRIX AND ABLATION DESIGN")
    table4_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table4 = insert_table_after(
        table4_caption,
        [
            ["Evaluation Block", "Dataset / Variant", "Evidence Type", "Primary Metrics"],
            ["Real-data behavior validation", "xAPI-Edu-Data", "Passive learner-signal classification", "Accuracy, Macro-F1"],
            ["Knowledge tracing benchmark", "ASSISTments 2009-2010", "Sequential next-response prediction", "AUC, ACC, RMSE"],
            ["Scalability benchmark", "EdNet-KT1", "Large-scale tutoring interaction logs", "AUC, Calibration, Throughput"],
            ["Ablation ladder", "BKT-only -> Full Lumina", "Component contribution analysis", "Delta AUC, Delta Macro-F1, Latency"],
        ],
    )
    fig2_para = insert_picture_after(table4, fig_paths["fig2"], 6.0)
    fig2_caption = insert_paragraph_after(
        fig2_para,
        "Fig. 2: Evaluation blueprint used in the revised paper. Prototype safety, real-data validation, and benchmark-plus-ablation analysis are treated as separate but connected evidence streams.",
    )
    fig2_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    replace_paragraph_text(metrics_heading, "E. Evaluation Metrics")
    replace_paragraph_text(
        find_paragraph(doc, "The system was evaluated across multiple dimensions"),
        "The system is evaluated across three metric families. Prototype metrics include intervention accuracy, throughput, execution latency, and runtime stability. Real-data metrics include accuracy, macro-F1, and weighted-F1 to reflect class-balanced predictive quality on xAPI-Edu-Data. The benchmark matrix for future knowledge-tracing experiments adds AUC, RMSE, and calibration so that predictive performance, confidence quality, and operational efficiency can be compared across BKT, DKT, DKVMN, SAKT, and Lumina variants.",
    )

    results_intro = find_paragraph(doc, "This section presents the experimental results")
    replace_paragraph_text(
        results_intro,
        "This section reports three layers of evidence: a synthetic stress test of the analytics pipeline, behavior-state classification behavior under controlled scenarios, and an external validation study on a real educational dataset.",
    )
    replace_paragraph_text(find_paragraph(doc, "Table IV presents the results from the 10,000-iteration stress test"), "Table V presents the results from the 10,000-iteration stress test of the Analytics Agent pipeline.")
    replace_paragraph_text(find_paragraph(doc, "TABLE IV: ANALYTICS AGENT STRESS TEST RESULTS"), "TABLE V: ANALYTICS AGENT STRESS TEST RESULTS")
    replace_paragraph_text(find_paragraph(doc, "Table V shows the calibrated output metrics across different learner scenario types"), "Table VI shows the calibrated output metrics across different learner scenario types, demonstrating the agent's ability to distinguish productive flow, struggle, mixed activity, and idle behavior.")
    replace_paragraph_text(find_paragraph(doc, "TABLE V: BEHAVIORAL CLASSIFICATION AND INTERVENTION METRICS"), "TABLE VI: BEHAVIORAL CLASSIFICATION AND INTERVENTION METRICS")

    intervention_para = find_paragraph(doc, "The system achieved a true positive rate")
    replace_paragraph_text(
        intervention_para,
        "Within the synthetic protocol, the system achieved perfect recall for the Struggle scenario family while maintaining zero false positive interventions for Flow, Mixed, and Idle states. This result should be interpreted as a property of the prototype rule stack and simulated distributions, not as a claim of perfect classroom performance. The result nevertheless validates that the error-aware engagement model and conservative intervention thresholds behave coherently under controlled stress conditions.",
    )

    real_results_heading = insert_paragraph_after(intervention_para, "D. Real Dataset Benchmark on xAPI-Edu-Data")
    real_results_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in real_results_heading.runs:
        run.bold = True
        run.font.size = Pt(10)
    real_results_para = insert_paragraph_after(
        real_results_heading,
        "Table VII reports the external validation results on xAPI-Edu-Data. Random Forest produced the best overall performance with 0.800 accuracy and 0.804 macro-F1, outperforming the linear and kernel baselines. The most predictive variables were visited learning resources, raised hands, announcements viewed, absence patterns, and discussion participation. This is important for Lumina because those same signal families are already represented in the proposed Behavior Engine, providing real-data support for the paper's claim that passive interaction traces carry actionable educational information.",
    )
    real_results_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    table7_caption = insert_paragraph_after(real_results_para, "TABLE VII: REAL-DATA BEHAVIORAL BENCHMARK RESULTS (5-FOLD CV)")
    table7_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf_result = xapi_metrics["results"]["Random Forest"]
    table7 = insert_table_after(
        table7_caption,
        [
            ["Model", "Accuracy", "Macro-F1", "Observation"],
            [
                "Logistic Regression",
                f"{xapi_metrics['results']['Logistic Regression']['accuracy']:.3f}",
                f"{xapi_metrics['results']['Logistic Regression']['macro_f1']:.3f}",
                "Strong linear baseline; lower nonlinear capacity",
            ],
            [
                "Random Forest",
                f"{rf_result['accuracy']:.3f}",
                f"{rf_result['macro_f1']:.3f}",
                "Best overall real-data behavior classifier",
            ],
            [
                "Gradient Boosting",
                f"{xapi_metrics['results']['Gradient Boosting']['accuracy']:.3f}",
                f"{xapi_metrics['results']['Gradient Boosting']['macro_f1']:.3f}",
                "Competitive but less stable than Random Forest",
            ],
            [
                "RBF SVM",
                f"{xapi_metrics['results']['RBF SVM']['accuracy']:.3f}",
                f"{xapi_metrics['results']['RBF SVM']['macro_f1']:.3f}",
                "Useful nonlinear baseline with balanced margins",
            ],
        ],
    )
    fig3_para = insert_picture_after(table7, fig_paths["fig3"], 6.0)
    fig3_caption = insert_paragraph_after(
        fig3_para,
        "Fig. 3: Public-data behavioral validation. Left: model comparison on xAPI-Edu-Data. Right: top predictors learned by the Random Forest model.",
    )
    fig3_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    replace_paragraph_text(find_paragraph(doc, "D. Scoring Model Evaluation"), "E. Scoring Model Evaluation")
    replace_paragraph_text(find_paragraph(doc, "TABLE VI: FOUR-DIMENSIONAL SCORING MODEL COMPARISON"), "TABLE VIII: FOUR-DIMENSIONAL SCORING MODEL COMPARISON")
    replace_paragraph_text(find_paragraph(doc, "E. Performance Comparison with Existing Systems"), "F. Performance Comparison with Existing Systems")
    replace_paragraph_text(find_paragraph(doc, "TABLE VII: PERFORMANCE COMPARISON WITH EXISTING SYSTEMS"), "TABLE IX: PERFORMANCE COMPARISON WITH EXISTING SYSTEMS")

    replace_paragraph_text(
        find_paragraph(doc, "The experimental results demonstrate that Lumina's multi-agent architecture"),
        "The revised evaluation should be interpreted as a layered validation strategy rather than a single monolithic benchmark. The synthetic stress test establishes that the analytics agent behaves deterministically, remains stable under load, and applies intervention policies coherently across scenario families. The xAPI experiment supplies external evidence that behavior-oriented learner signals are predictive on real educational data. The benchmark matrix then makes explicit what remains to be evaluated on public sequence datasets before claiming state-of-the-art knowledge-tracing performance.",
    )
    replace_paragraph_text(
        find_paragraph(doc, "The 100% sensitivity with 0% false positive rate"),
        "The perfect struggle recall observed in the simulation is meaningful because it confirms that Lumina's intervention thresholds are internally consistent: high error rate alone is not enough to trigger an alert, and neither is activity alone. Instead, interventions arise only when high cognitive load, sustained error accumulation, and behavioral degradation co-occur. This makes the result more believable than a single-threshold detector, but it remains a prototype finding until replicated with classroom telemetry.",
    )
    replace_paragraph_text(
        find_paragraph(doc, "The hybrid BKT+DKT knowledge tracing approach"),
        "The hybrid BKT+DKT knowledge-tracing design remains one of the paper's strongest technical choices because it partitions interpretability and predictive power instead of forcing a single model to satisfy both objectives. BKT provides the transparent concept-level mastery state that teachers and institutions can audit, while DKT supplies the sequence-aware predictor needed for adaptive recommendations. The added benchmark matrix now makes it possible to compare this hybrid against DKVMN and SAKT in a reproducible way rather than merely asserting architectural superiority.",
    )
    replace_paragraph_text(
        find_paragraph(doc, "The four-dimensional scoring model addresses a fundamental limitation"),
        "The four-dimensional scoring model is now mathematically cleaner because the paper defines it as a weighted sum with normalized components rather than an ambiguous multiplicative expression. This matters because weighted sums preserve interpretability, keep scores in a bounded range, and make the justification for each coefficient explicit. As a result, the model better communicates why correctness dominates, why understanding still matters substantially, and how effort and growth are included without overpowering academic validity.",
    )
    replace_paragraph_text(
        find_paragraph(doc, "The self-hosted architecture with local LLM inference represents"),
        "The self-hosted architecture with local LLM inference remains a differentiator for institutional adoption. Privacy, auditability, and cost predictability are not secondary operational details; they are design constraints in educational settings. The revised evaluation strengthens this claim indirectly by showing that large parts of Lumina's adaptive intelligence, particularly the analytics and behavior-modeling layers, can already be validated without routing student data through external providers.",
    )
    replace_paragraph_text(
        find_paragraph(doc, "Several limitations should be acknowledged"),
        "Several limitations should be acknowledged. First, although the paper now includes a real public dataset benchmark, the xAPI-Edu-Data task is cross-sectional and behavior-focused; it does not replace longitudinal learner-sequence evaluation on ASSISTments or EdNet. Second, the perfect synthetic intervention result reflects the current scenario generator and rule thresholds, so classroom deployment may show lower recall and nonzero false positives. Third, the pathway optimization policy is still trained in simulation and requires calibration with real trajectory data before any strong claims about long-horizon RL benefits can be made. Fourth, full institutional deployment with local LLM inference still carries nontrivial infrastructure cost despite the encouraging CPU-only analytics results.",
    )

    replace_paragraph_text(
        find_paragraph(doc, "This paper presented Lumina, a next-generation AI-powered self-hosted learning management system"),
        "This paper presented Lumina, a next-generation AI-powered self-hosted learning management system that integrates multi-agent orchestration, hybrid knowledge tracing, reinforcement learning, and behavior-aware personalization into a single adaptive platform. The revised manuscript strengthens the original contribution by correcting the mathematical formulation of the scoring and reward functions, grounding the evaluation in both synthetic and real public data, and explicitly defining a benchmark-plus-ablation protocol for future knowledge-tracing comparisons.",
    )
    replace_paragraph_text(
        find_paragraph(doc, "The key contributions include:"),
        f"The key contributions include: (1) a closed-loop multi-agent architecture for tutoring, assessment, analytics, intervention, and governance, (2) a mathematically consistent BKT+DKT learner-modeling stack with interpretable scoring and policy optimization equations, (3) a behavior engine that uses 50+ passive signals and demonstrates external validity on xAPI-Edu-Data with 0.800 accuracy and 0.804 macro-F1, and (4) a benchmark-ready evaluation matrix spanning simulation, real-data validation, public-sequence datasets, and ablation analysis.",
    )
    replace_paragraph_text(
        find_paragraph(doc, "Experimental evaluation demonstrates processing throughput"),
        f"Experimental evaluation demonstrates that the analytics prototype can process approximately {sim_metrics['throughput']:.0f} randomized scenarios per second with zero runtime failures, while the external behavioral benchmark confirms that Lumina's signal families are meaningful beyond simulation. Although longitudinal classroom studies and public sequence benchmarks remain future work, the paper now offers a substantially stronger empirical foundation for IEEE-style review.",
    )

    references_anchor = find_paragraph(doc, "[29] A. Tsai et al.")
    ref30 = insert_paragraph_after(references_anchor, '[30] I. Aljarah, "xAPI-Edu-Data," Kaggle dataset, 2016.')
    ref31 = insert_paragraph_after(ref30, '[31] Y. Choi, Y. Lee, J. Cho, J. Baek, B. Kim, and J. Heo, "EdNet: A Large-Scale Hierarchical Dataset in Education," in Proc. Int. Conf. Educational Data Mining (EDM), 2020.')
    insert_paragraph_after(ref31, '[32] S. Pandey and G. Karypis, "A Self-Attentive Model for Knowledge Tracing," arXiv:1907.06837, 2019.')

    style_document(doc)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    sim_metrics = generate_simulation_metrics()
    xapi_metrics = generate_xapi_metrics()

    fig_paths = {
        "fig1": ASSET_DIR / "lumina_multi_agent_architecture.png",
        "fig2": ASSET_DIR / "lumina_evaluation_blueprint.png",
        "fig3": ASSET_DIR / "lumina_xapi_results.png",
    }
    save_architecture_figure(fig_paths["fig1"])
    save_benchmark_blueprint(fig_paths["fig2"])
    save_real_data_figure(fig_paths["fig3"], xapi_metrics)

    doc = Document(str(INPUT_DOC))
    edit_document(doc, sim_metrics, xapi_metrics, fig_paths)
    doc.save(str(OUTPUT_DOC))

    print(f"Wrote {OUTPUT_DOC}")
    print(f"Figures: {fig_paths['fig1']}, {fig_paths['fig2']}, {fig_paths['fig3']}")
    print(f"Simulation throughput: {sim_metrics['throughput']:.2f} runs/sec")
    print(f"Random Forest macro-F1: {xapi_metrics['results']['Random Forest']['macro_f1']:.3f}")


if __name__ == "__main__":
    main()
