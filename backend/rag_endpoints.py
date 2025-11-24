"""
RAG-specific API endpoints for Lumina LMS
Handles content ingestion, search, and LLM-powered query answering
"""

from fastapi import HTTPException, BackgroundTasks, UploadFile, File, Depends, APIRouter
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import uvicorn
import logging
import uuid
import os
import jwt
from datetime import datetime
import json

from services.embeddings import EmbeddingService
from services.vector_store import VectorStoreService
from services.content_parser import ContentParser
from services.llm_service import LLMService
from config import settings

# Initialize services
embedding_service = EmbeddingService()
vector_store = VectorStoreService()
content_parser = ContentParser()
llm_service = LLMService()

from config import settings
from models_rag import (
    ContentIngestionRequest, SearchRequest, QueryRequest,
    FileUploadMetadata, ProcessingStatus, SearchResult, QueryResponse
)

# Create router
router = APIRouter(prefix="/rag", tags=["RAG"])

# JWT Security
security = HTTPBearer()

# Configure logging
logger = logging.getLogger(__name__)

# Pydantic models for RAG endpoints
class QueryRequest(BaseModel):
    query: str
    course_id: Optional[str] = None
    limit: int = 5

# JWT authentication dependency
async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Verify JWT token and return user info"""
    try:
        payload = jwt.decode(credentials.credentials, settings.SECRET_KEY, algorithms=[settings.JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

async def process_file_async(doc_id: str, file_path: str, content_type: str, metadata: Dict[str, Any]):
    """Process uploaded file for RAG system"""
    try:
        logger.info(f"Processing file asynchronously: {doc_id}")

        # Parse file content based on type
        if content_type == 'text/plain':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
        elif content_type == 'application/pdf':
            try:
                from pdfminer.high_level import extract_text
                content = extract_text(file_path)
            except ImportError:
                logger.error("PDF parsing not available - pdfminer not installed")
                raise
        elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            except ImportError:
                logger.error("DOCX parsing not available - python-docx not installed")
                raise
        else:
            raise ValueError("Unsupported file type")

        # Chunk content using ContentParser service
        chunks = content_parser.parse_and_chunk(
            content,
            chunk_size=1000,
            chunk_overlap=200
        )

        # Prepare chunks for batch processing
        chunks_data = []
        for i, chunk in enumerate(chunks):
            # Generate embedding
            embedding = embedding_service.generate_embedding(chunk)
            
            # Prepare chunk data
            chunk_id = f"{doc_id}_chunk_{i}"
            chunk_metadata = {
                **metadata,
                "chunk_num": i,
                "total_chunks": len(chunks),
                "document_id": doc_id,
                "chunk_id": chunk_id
            }
            
            chunks_data.append({
                "id": chunk_id,
                "content": chunk,
                "embedding": embedding,
                "metadata": chunk_metadata
            })

        # Batch insert into vector store
        success = await vector_store.insert_batch(
            ids=[c["id"] for c in chunks_data],
            contents=[c["content"] for c in chunks_data],
            embeddings=[c["embedding"] for c in chunks_data],
            metadatas=[c["metadata"] for c in chunks_data]
        )

        if not success:
            raise Exception("Failed to insert chunks into vector store")

        # Update document status in database
        from db import get_db
        from sqlalchemy.orm import Session
        db: Session = next(get_db())
        db.execute(
            """
            UPDATE documents 
            SET status = 'processed', 
                processed_chunks = :chunk_count,
                updated_at = NOW()
            WHERE id = :doc_id
            """,
            {"doc_id": doc_id, "chunk_count": len(chunks)}
        )
        db.commit()

        # Clean up temp file
        try:
            os.remove(file_path)
        except Exception as e:
            logger.warning(f"Failed to remove temp file {file_path}: {str(e)}")

        logger.info(f"File processed successfully: {doc_id}, chunks: {len(chunks)}")

    except Exception as e:
        logger.error(f"File processing error for {doc_id}: {str(e)}")
        # Update document status to error
        try:
            db: Session = next(get_db())
            db.execute(
                """
                UPDATE documents 
                SET status = 'error',
                    error_message = :error,
                    updated_at = NOW()
                WHERE id = :doc_id
                """,
                {"doc_id": doc_id, "error": str(e)}
            )
            db.commit()
        except Exception as db_err:
            logger.error(f"Failed to update error status: {str(db_err)}")
        raise

    except Exception as e:
        logger.error(f"File processing error: {str(e)}")

# File upload endpoint for RAG content processing
@router.post("/upload", response_model=ProcessingStatus)
async def upload_file(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    metadata: FileUploadMetadata = Depends(),
    current_user: dict = Depends(get_current_user)
):
    """Upload and process files for RAG system"""
    try:
        logger.info(f"Uploading file: {file.filename}")

        # Validate file type
        allowed_types = ['text/plain', 'application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']
        if file.content_type not in allowed_types:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        # Read file content
        file_content = await file.read()
        file_size = len(file_content)

        # Generate unique filename and document ID
        doc_id = str(uuid.uuid4())
        filename = f"{doc_id}_{file.filename}"
        temp_path = f"/tmp/{filename}"

        # Save file temporarily
        with open(temp_path, "wb") as f:
            f.write(file_content)

        # Store document metadata
        from db import get_db
        from models_rag import Document
        from sqlalchemy.orm import Session

        db: Session = next(get_db())

        document = Document(
            id=doc_id,
            filename=filename,
            original_filename=file.filename,
            content_type=file.content_type,
            file_size=file_size,
            uploaded_by=current_user.get('id'),
            course_id=metadata.course_id,
            title=metadata.title or file.filename,
            description=metadata.description,
            metadata={
                "content_type": metadata.content_type,
                "tags": metadata.tags,
                "upload_date": datetime.utcnow().isoformat()
            }
        )

        db.add(document)
        db.commit()

        # Start background processing
        background_tasks.add_task(
            process_file_async,
            doc_id=doc_id,
            file_path=temp_path,
            content_type=file.content_type,
            metadata=document.metadata
        )

        return ProcessingStatus(
            status="processing",
            message="File upload successful, processing started",
            document_id=doc_id
        )

    except Exception as e:
        logger.error(f"File upload error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# RAG endpoints
@router.post("/ingest", response_model=ProcessingStatus)
async def ingest_content(
    content: ContentIngestionRequest,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Ingest content into the vector store"""
    try:
        # Parse and chunk content
        chunks = content_parser.parse_and_chunk(
            content.text,
            chunk_size=content.chunk_size,
            chunk_overlap=content.chunk_overlap
        )
        
        # Process chunks and prepare for batch insertion
        chunks_data = []
        for i, chunk in enumerate(chunks):
            embedding = embedding_service.generate_embedding(chunk)
            chunks_data.append({
                "id": str(uuid.uuid4()),
                "content": chunk,
                "embedding": embedding,
                "metadata": {
                    **content.metadata,
                    "course_id": content.course_id,
                    "lesson_id": content.lesson_id,
                    "chunk_num": i,
                    "source": content.source,
                    "type": content.content_type,
                    "title": content.title,
                    "timestamp": datetime.utcnow().isoformat()
                }
            })
        
        # Schedule background batch insertion
        background_tasks.add_task(
            vector_store.insert_batch,
            ids=[c["id"] for c in chunks_data],
            contents=[c["content"] for c in chunks_data],
            embeddings=[c["embedding"] for c in chunks_data],
            metadatas=[c["metadata"] for c in chunks_data]
        )
        
        return ProcessingStatus(
            status="processing",
            message="Content ingestion started",
            total_chunks=len(chunks),
            processed_chunks=0
        )

    except Exception as e:
        logger.error(f"Content ingestion error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/search", response_model=List[SearchResult])
async def semantic_search(
    query: SearchRequest,
    current_user: dict = Depends(get_current_user)
):
    """Search for relevant content with optional filters"""
    try:
        # Generate query embedding
        query_embedding = embedding_service.generate_embedding(query.text)
        
        # Prepare filters
        filters = {
            "course_id": query.course_id,
            "lesson_id": query.lesson_id
        } if query.course_id or query.lesson_id else None
        
        # Perform vector search
        results = await vector_store.search(
            query_embedding=query_embedding,
            filter_dict=filters,
            limit=query.limit
        )
        
        return [SearchResult(**result) for result in results]

    except Exception as e:
        logger.error(f"Search error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/query", response_model=QueryResponse)
async def query_content(
    request: QueryRequest,
    current_user: dict = Depends(get_current_user)
):
    """Query the RAG system for content retrieval and generation"""
    try:
        start_time = datetime.utcnow()
        logger.info(f"Processing query: {request.query[:50]}...")

        # Generate query embedding
        query_embedding = embedding_service.generate_embedding(request.query)

        # Prepare filters
        filters = {
            "course_id": request.course_id,
            "lesson_id": request.lesson_id
        } if request.course_id or request.lesson_id else None

        # Search for relevant chunks
        search_results = await vector_store.search(
            query_embedding=query_embedding,
            filter_dict=filters,
            limit=request.context_limit
        )

        if not search_results:
            return QueryResponse(
                query=request.query,
                answer="I couldn't find any relevant information to answer your question.",
                sources=[],
                confidence=0.0,
                processing_time=(datetime.utcnow() - start_time).total_seconds()
            )

        # Prepare context for LLM
        context_chunks = []
        for result in search_results:
            chunk = f"Content: {result['content']}\n"
            if result['metadata'].get('title'):
                chunk = f"Title: {result['metadata']['title']}\n" + chunk
            context_chunks.append(chunk)

        context = "\n---\n".join(context_chunks)

        # Generate prompt
        prompt = f"""Answer the question using ONLY the provided context. If the context doesn't contain enough information to answer confidently, say so.

Context:
{context}

Question: {request.query}

Answer: """

        # Generate response using LLM
        llm_response = await llm_service.generate(
            prompt=prompt,
            max_tokens=request.max_tokens,
            temperature=0.7
        )

        # Calculate confidence based on search scores
        avg_score = sum(r['score'] for r in search_results) / len(search_results)
        confidence = min(1.0, max(0.0, avg_score))  # Normalize to 0-1

        processing_time = (datetime.utcnow() - start_time).total_seconds()

        return QueryResponse(
            query=request.query,
            answer=llm_response.get("result", "I couldn't generate a response."),
            sources=[SearchResult(**result) for result in search_results],
            confidence=confidence,
            processing_time=processing_time
        )

    except Exception as e:
        logger.error(f"Query error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
